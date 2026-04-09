#!/usr/bin/env python3
"""Generate Client_Engagement_Models.docx (design-phase PM doc)."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parent
DIAG = ROOT / "_diagrams"
OUT = ROOT / "Client_Engagement_Models.docx"


def add_caption(doc, text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    r.bold = bold
    r.font.size = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def main():
    doc = Document()

    t = doc.add_heading("Client engagement models", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("LexDesk / Advocate CRM — product design note")
    meta = doc.add_paragraph()
    meta.add_run("Phase: ").bold = True
    meta.add_run("Design (pre-development). ")
    meta.add_run("Audience: ").bold = True
    meta.add_run("Product, design, leadership.")

    doc.add_heading("1. Purpose", level=1)
    doc.add_paragraph(
        "This document captures agreed client-side engagement rules: how a client "
        "enters the platform, whom they connect with, and how matters and messaging behave "
        "for a pure solo advocate, for an advocate who practises independently on the "
        "platform while employed or affiliated with a firm, and for a firm-led model where "
        "assignment and reassignment of advocates is controlled by the firm."
    )

    doc.add_heading("2. Global product rules (fixed)", level=1)
    doc.add_paragraph(
        "Geography: India-wide framing; court and matter metadata must support all states even if integrations roll out incrementally.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Messaging: chat is available only between linked clients and the appropriate advocate-side parties (no open messaging from directory).",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Dual-sided product: advocate CRM and client app are co-equal; permissions and matter context must be explicit in each flow.",
        style="List Bullet",
    )

    doc.add_heading("3. Three client selection approaches", level=1)

    doc.add_heading("3.1 Path 1 — Pure individual advocate", level=2)
    doc.add_paragraph(
        "The advocate operates as a solo practice on the platform. On the client app, the "
        "client discovers that advocate, sends a connection request, and after acceptance "
        "sees that same advocate as the face of every matter. There is no firm orchestration "
        "layer visible to the client."
    )

    doc.add_heading("3.2 Path 2 — Advocate who belongs to a firm but uses the platform independently", level=2)
    doc.add_paragraph(
        "The advocate is part of a firm in real life (and may appear on a firm roster in "
        "CRM), but on the client experience they opt into being contactable as an individual. "
        "The client connects to that person; matters and updates are anchored to that "
        "advocate. The platform does not, for that relationship, apply firm-driven "
        "reassignment rules unless the product later migrates that client to Path 3 with consent."
    )

    doc.add_heading("3.3 Path 3 — Firm-led engagement", level=2)
    doc.add_paragraph(
        "The client engages with the firm (or with a firm-designated lead such as a managing "
        "partner). The firm accepts the client, opens matters under the firm, and assigns a "
        "lead advocate per matter. The firm may change the assigned advocate later; the client "
        "app should show firm identity together with the current matter lead so expectations stay clear."
    )

    doc.add_heading("4. Client-side display summary", level=1)

    table = doc.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Path"
    hdr[1].text = "Whom the client primarily sees"
    hdr[2].text = "Who controls advocate assignment"

    row1 = table.rows[1].cells
    row1[0].text = "1 — Pure individual"
    row1[1].text = "Single advocate only"
    row1[2].text = "N/A (same advocate)"

    row2 = table.rows[2].cells
    row2[0].text = "2 — Firm member, independent on platform"
    row2[1].text = "The chosen advocate (person-first)"
    row2[2].text = "Effectively the advocate for that linked relationship; firm does not reassign on platform for that path"

    row3 = table.rows[3].cells
    row3[0].text = "3 — Firm-led"
    row3[1].text = "Firm name plus lead advocate per matter"
    row3[2].text = "Firm (assignment and reassignment)"

    doc.add_paragraph()

    doc.add_heading("5. Flow diagrams", level=1)
    doc.add_paragraph("High-level flows for design and stakeholder review. Not a technical specification.")

    diagrams = [
        ("overview_three_paths.png", "Figure 1 — Overview: three client entry paths"),
        ("path1_pure_individual.png", "Figure 2 — Path 1: pure individual advocate"),
        ("path2_firm_member_independent.png", "Figure 3 — Path 2: firm member, independent client relationship on platform"),
        ("path3_firm_led.png", "Figure 4 — Path 3: firm-led assignment and reassignment"),
    ]

    for filename, caption in diagrams:
        path = DIAG / filename
        if not path.exists():
            raise SystemExit(f"Missing diagram: {path}")
        doc.add_paragraph()
        doc.add_picture(str(path), width=Inches(6.0))
        add_caption(doc, caption)

    doc.add_heading("6. Open design follow-ups (optional)", level=1)
    doc.add_paragraph(
        "Connection acceptance: whether Path 3 requests are accepted only by firm admin, "
        "only by matter lead, or by configurable firm policy.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Migration: rules if a Path 2 client later moves under firm-led Path 3.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Directory: whether firm-branded listings are required alongside advocate profiles.",
        style="List Bullet",
    )

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
